from flet import *
from datetime import datetime, timezone
import logging
import asyncio
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import io
import os
import base64

# Configurar logging
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)

class TutorDarkMoodPalette:
    MAIN_BACKGROUND = "#1E252D"
    CARD_SURFACE = "#2A323C"
    BORDER_SUBTLE = "#3B4653"
    TRUST = "#90CAF9"            # Pastel blue for depression (BDI)
    TRUST_HOVER = "#B3E0FF"      # Lighter pastel blue for hover
    ALERT = "#FFCC80"            # Pastel yellow for anxiety (BAI)
    SUPPORT = "#5C9CE6"
    URGENT = "#FFAB91"           # Pastel red for stress (PSS)
    TEXT_MAIN = "#E2E8F0"
    TEXT_ON_ACCENT = "#FFFFFF"
    SUCCESS_FEEDBACK = "#2DD4BF"
    ERROR_FEEDBACK = "#F87171"

class FilterContent(Column):
    def __init__(self, page, tutor_data=None, selected_group=None, on_group_change=None, cache=None):
        super().__init__()
        self.page = page
        self.tutor_data = tutor_data or {}
        self.on_group_change = on_group_change
        self.cache = cache
        self.selected_group = selected_group
        self.year = datetime.now().year
        self.selected_cuatrimestre = f"Todo {self.year}"
        self.selected_student = None
        # Dropdown de grupos
        self.group_dropdown = Dropdown(
            options=[],
            width=200,
            bgcolor=TutorDarkMoodPalette.CARD_SURFACE,
            border_color=TutorDarkMoodPalette.BORDER_SUBTLE,
            color=TutorDarkMoodPalette.TEXT_MAIN,
            on_change=lambda e: self.page.run_task(self.update_group, e.control.value)
        )
        # Botón de reporte
        self.report_button = IconButton(
            icon=Icons.DESCRIPTION,
            icon_size=20,
            icon_color=TutorDarkMoodPalette.TEXT_ON_ACCENT,
            bgcolor=TutorDarkMoodPalette.TRUST,
            width=48,
            height=48,
            tooltip="Generar Reporte en Word",
            style=ButtonStyle(
                shape=RoundedRectangleBorder(radius=8),
                elevation=2,
                overlay_color=Colors.with_opacity(0.2, TutorDarkMoodPalette.TRUST_HOVER)
            ),
            on_click=self.generate_report
        )
        # Dropdown de cuatrimestre dinámico
        self.cuatrimestre_dropdown = Dropdown(
            value=self.selected_cuatrimestre,
            options=[
                dropdown.Option(f"Ene-Abr {self.year}"),
                dropdown.Option(f"May-Ago {self.year}"),
                dropdown.Option(f"Sep-Dic {self.year}"),
                dropdown.Option(f"Todo {self.year}")
            ],
            width=200,
            bgcolor=TutorDarkMoodPalette.CARD_SURFACE,
            border_color=TutorDarkMoodPalette.BORDER_SUBTLE,
            color=TutorDarkMoodPalette.TEXT_MAIN,
            on_change=self.on_cuatrimestre_change
        )
        # Contenedor lista alumnos
        self.student_list_container = Column(
            spacing=4,
            scroll="auto",
            height=400,  # Fixed height to match chart container
            width=300   # Fixed width to maintain original size
        )
        # Contenedor info alumno
        self.student_info_container = Container(
            content=Column([Text("Selecciona un alumno", color=TutorDarkMoodPalette.TEXT_MAIN)], scroll="auto"),
            bgcolor=TutorDarkMoodPalette.CARD_SURFACE,
            padding=padding.all(16),
            border_radius=8,
            border=border.all(1, TutorDarkMoodPalette.BORDER_SUBTLE),
            height=200,
            expand=True
        )
        # Contenedor recomendaciones
        self.recommendations_container = Container(
            content=Column([
                Text("Recomendaciones", size=16, weight="bold"),
                Text("Ansiedad: N/A", color=TutorDarkMoodPalette.TEXT_MAIN),
                Text("Depresión: N/A", color=TutorDarkMoodPalette.TEXT_MAIN),
                Text("Estrés: N/A", color=TutorDarkMoodPalette.TEXT_MAIN)
            ], spacing=8, scroll="auto"),
            bgcolor=TutorDarkMoodPalette.CARD_SURFACE,
            padding=padding.all(16),
            border_radius=8,
            border=border.all(1, TutorDarkMoodPalette.BORDER_SUBTLE),
            height=200,
            expand=True
        )
        # Contenedor gráfica
        self.chart_container = Container(
            content=Text("Selecciona un alumno para ver su historial",
                        color=TutorDarkMoodPalette.TEXT_MAIN),
            bgcolor=TutorDarkMoodPalette.CARD_SURFACE,
            padding=padding.all(16),
            border_radius=8,
            border=border.all(1, TutorDarkMoodPalette.BORDER_SUBTLE),
            height=400,
            expand=True
        )
        # Contenedor sección historial
        self.historial_container = Container(
            content=Column([
                Row([
                    Text("Historial", size=16, color=TutorDarkMoodPalette.TEXT_MAIN),
                    self.cuatrimestre_dropdown
                ], alignment="spaceBetween"),
                self.chart_container
            ], expand=True),
            bgcolor=TutorDarkMoodPalette.CARD_SURFACE,
            padding=padding.all(8),
            border_radius=8,
            border=border.all(1, TutorDarkMoodPalette.BORDER_SUBTLE),
            height=450,
            expand=True
        )
        # Layout principal
        self.controls = [
            Row([
                Text("Filtrado de alumnos", size=20, color=TutorDarkMoodPalette.TEXT_MAIN, expand=True),
                self.group_dropdown,
                self.report_button
            ], alignment="spaceBetween"),
            Row([
                self.student_info_container,
                self.recommendations_container
            ], spacing=10, expand=True),
            Row([
                self.historial_container,
                Container(
                    content=self.student_list_container,
                    bgcolor=TutorDarkMoodPalette.CARD_SURFACE,
                    border=border.all(1, TutorDarkMoodPalette.BORDER_SUBTLE),
                    border_radius=8,
                    padding=padding.all(8),
                    width=300,  # Match student_list_container width
                    height=450  # Match historial_container height
                )
            ], spacing=10, expand=True)
        ]
        self.bgcolor = TutorDarkMoodPalette.MAIN_BACKGROUND
        self.expand = True
        self.spacing = 10
        self.padding = padding.all(10)

    def show_snackbar(self, message, color):
        self.page.snack_bar = SnackBar(
            content=Text(
                message,
                color=TutorDarkMoodPalette.TEXT_ON_ACCENT,
                font_family="Fredoka"
            ),
            behavior=SnackBarBehavior.FLOATING,
            shape=RoundedRectangleBorder(radius=8),
            bgcolor=color,
            elevation=8,
            margin=padding.all(20),
            show_close_icon=True
        )
        self.page.snack_bar.open = True
        self.page.update()

    async def initialize(self):
        tutor_groups = self.tutor_data.get("groups", [])
        if not tutor_groups:
            self.show_no_groups_message()
            return
        self.group_dropdown.options = [dropdown.Option(g) for g in tutor_groups]
        self.selected_group = tutor_groups[0] if not self.selected_group else self.selected_group
        self.group_dropdown.value = self.selected_group
        await self.update_group(self.selected_group)
        if self.page:
            self.page.update()

    def show_no_groups_message(self):
        message = Column([
            Text("No tienes grupos asignados", size=18, weight="bold"),
            Text("Contacta al administrador para que te asigne grupos", size=14)
        ], alignment="center", horizontal_alignment="center", expand=True)
        self.controls = [
            Container(
                content=message,
                bgcolor=TutorDarkMoodPalette.CARD_SURFACE,
                padding=40,
                border_radius=8,
                expand=True
            )
        ]

    async def update_group(self, new_group):
        self.selected_group = new_group
        self.selected_student = None
        students = self.cache.get_users_by_group(new_group) if self.cache else []
        await self.build_student_list(students)
        if students:
            await self.select_student(students[0])
        else:
            self.student_info_container.content = Column([Text(
                f"No hay alumnos en el grupo {new_group}",
                color=TutorDarkMoodPalette.TEXT_MAIN
            )], scroll="auto")
            self.recommendations_container.content = Column([Text(
                "Selecciona un alumno para ver recomendaciones",
                color=TutorDarkMoodPalette.TEXT_MAIN
            )], scroll="auto")
            self.chart_container.content = Text(
                "No hay datos para mostrar",
                color=TutorDarkMoodPalette.TEXT_MAIN
            )
        if self.page:
            self.page.update()

    async def build_student_list(self, students):
        self.student_list_container.controls.clear()
        if not students:
            self.student_list_container.controls.append(
                Text("No hay alumnos en este grupo",
                    color=TutorDarkMoodPalette.TEXT_MAIN)
            )
            return
        for student in students:
            self.student_list_container.controls.append(
                ListTile(
                    title=Text(student.get("name", "Nombre no disponible"),
                    color=TutorDarkMoodPalette.TEXT_MAIN),
                    subtitle=Text(student.get("student_id", ""),
                                 color=TutorDarkMoodPalette.TEXT_MAIN, size=12),
                    on_click=lambda e, s=student: self.page.run_task(self.select_student, s)
                )
            )

    async def select_student(self, student):
        self.selected_student = student
        self.update_student_info(student)
        self.update_recommendations(student)
        self.chart_container.content = await self.create_chart(student.get('doc_id', ''), self.selected_cuatrimestre)
        if self.page:
            self.page.update()

    def update_student_info(self, student):
        last_login = student.get('lastLogin', 'N/A')
        if isinstance(last_login, datetime):
            last_login = last_login.strftime("%d %b")
        info = Column([
            Text(f"Nombre: {student.get('name', 'N/A')}", size=16, color=TutorDarkMoodPalette.TEXT_MAIN),
            Text(f"Edad: {student.get('age', 'N/A')}", color=TutorDarkMoodPalette.TEXT_MAIN),
            Text(f"Carrera: {student.get('class', 'N/A')}", color=TutorDarkMoodPalette.TEXT_MAIN),
            Text(f"Grupo: {student.get('group', 'N/A')}", color=TutorDarkMoodPalette.TEXT_MAIN),
            Text(f"Email: {student.get('email', 'N/A')}", color=TutorDarkMoodPalette.TEXT_MAIN),
            Text(f"Género: {student.get('gender', 'N/A')}", color=TutorDarkMoodPalette.TEXT_MAIN),
            Text(f"Activo: {'Sí' if student.get('isActive', False) else 'No'}", color=TutorDarkMoodPalette.TEXT_MAIN),
            Text(f"Último acceso: {last_login}", color=TutorDarkMoodPalette.TEXT_MAIN),
        ], spacing=6, scroll="auto")
        self.student_info_container.content = info

    def update_recommendations(self, student):
        recs = self.cache.get_user_recommendations(student.get('doc_id', '')) if self.cache else {}
        content = Column([
            Text("Recomendaciones actuales", size=16, weight="bold"),
            Text(f"Ansiedad: {recs.get('BAI', 'N/A')}", color=TutorDarkMoodPalette.TEXT_MAIN),
            Text(f"Depresión: {recs.get('BDI', 'N/A')}", color=TutorDarkMoodPalette.TEXT_MAIN),
            Text(f"Estrés: {recs.get('PSS', 'N/A')}", color=TutorDarkMoodPalette.TEXT_MAIN)
        ], spacing=8, scroll="auto")
        self.recommendations_container.content = content

    def get_month_name(self, month, year):
        month_names = {
            1: "Ene", 2: "Feb", 3: "Mar", 4: "Abr", 5: "May", 6: "Jun",
            7: "Jul", 8: "Ago", 9: "Sep", 10: "Oct", 11: "Nov", 12: "Dic"
        }
        return f"{month_names[month]} {year}"

    def month_to_number(self, month_name):
        month_map = {
            "Ene": 1, "Feb": 2, "Mar": 3, "Abr": 4, "May": 5, "Jun": 6,
            "Jul": 7, "Ago": 8, "Sep": 9, "Oct": 10, "Nov": 11, "Dic": 12
        }
        return month_map.get(month_name, 1)

    def get_quarter(self, date: datetime):
        month = date.month
        year = date.year
        if 1 <= month <= 4:
            return f"Ene-Abr {year}"
        elif 5 <= month <= 8:
            return f"May-Ago {year}"
        else:
            return f"Sep-Dic {year}"

    async def generate_report(self, e):
        try:
            group = self.group_dropdown.value
            if not group:
                self.show_snackbar("Selecciona un grupo", TutorDarkMoodPalette.ERROR_FEEDBACK)
                return
            users = self.cache.get_users_by_group(group) if self.cache else []
            if not users:
                self.show_snackbar(f"No hay alumnos en el grupo {group}", TutorDarkMoodPalette.ERROR_FEEDBACK)
                return

            doc = Document()
            # Cover page
            doc.add_heading("Reporte de Evaluaciones Psicológicas", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Grupo: {group}", style='Heading 2').alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d %b %Y %H:%M')}", style='Heading 3').alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("Generado por SERENIA Tutores", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

            for user in users:
                user_id = user.get("doc_id", "")
                user_name = user.get("name", "Sin nombre")
                # Student section
                doc.add_heading(f"Alumno: {user_name}", level=1)
                doc.add_paragraph(f"Email: {user.get('email', 'Sin email')}")
                doc.add_paragraph(f"Grupo: {user.get('group', 'Sin grupo')}")
                doc.add_paragraph(f"Edad: {user.get('age', 'N/A')}")
                doc.add_paragraph(f"Carrera: {user.get('class', 'N/A')}")
                doc.add_paragraph(f"Género: {user.get('gender', 'N/A')}")
                doc.add_paragraph(f"Activo: {'Sí' if user.get('isActive', False) else 'No'}")
                last_login = user.get('lastLogin', 'N/A')
                if isinstance(last_login, datetime):
                    last_login = last_login.strftime("%d %b")
                doc.add_paragraph(f"Último acceso: {last_login}")
                
                # Results table
                responses = self.cache.get_user_responses(user_id) if self.cache else []
                if responses:
                    doc.add_heading("Resultados de Cuestionarios", level=2)
                    table = doc.add_table(rows=1, cols=3)  # Reduced to 3 columns
                    table.style = 'Table Grid'
                    table.autofit = True
                    hdr_cells = table.rows[0].cells
                    headers = ['Fecha', 'Cuestionario', 'Nivel']
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = header
                        run = hdr_cells[i].paragraphs[0].runs[0]
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                        run.bold = True
                    for response in responses:
                        date = response.get("date")
                        if isinstance(date, str):
                            try:
                                date = datetime.fromisoformat(date.replace("Z", "+00:00"))
                            except ValueError:
                                date_str = str(date)
                            else:
                                date_str = date.strftime('%d %b')
                        else:
                            date_str = date.strftime('%d %b') if isinstance(date, datetime) else str(date)
                        row_cells = table.add_row().cells
                        row_cells[0].text = date_str
                        row_cells[1].text = response.get("questionnaire", "N/A")
                        row_cells[2].text = str(response.get("level", "N/A"))
                        for cell in row_cells:
                            run = cell.paragraphs[0].runs[0]
                            run.font.size = Pt(10)
                            run.font.name = 'Arial'

                # Chart
                quarters = {}
                for response in responses:
                    if isinstance(response["date"], str):
                        try:
                            date = datetime.fromisoformat(response["date"].replace("Z", "+00:00"))
                        except ValueError:
                            continue
                    else:
                        date = response["date"]
                    quarter = self.get_quarter(date)
                    questionnaire = response.get("questionnaire", "").strip().upper()
                    level = response.get("level", 0)
                    if quarter not in quarters:
                        quarters[quarter] = {"BAI": None, "BDI": None, "PSS": None}
                    quarters[quarter][questionnaire] = level

                quarter_list = sorted(quarters.keys())
                if quarter_list:
                    fig, ax = plt.subplots(figsize=(6, 4))
                    bai_levels = [quarters[q]["BAI"] if quarters[q]["BAI"] is not None else None for q in quarter_list]
                    bdi_levels = [quarters[q]["BDI"] if quarters[q]["BDI"] is not None else None for q in quarter_list]
                    pss_levels = [quarters[q]["PSS"] if quarters[q]["PSS"] is not None else None for q in quarter_list]
                    ax.plot(quarter_list, bai_levels, label="Ansiedad", color=TutorDarkMoodPalette.ALERT, marker='o', linestyle='-', linewidth=2)
                    ax.plot(quarter_list, bdi_levels, label="Depresión", color=TutorDarkMoodPalette.TRUST, marker='o', linestyle='-', linewidth=2)
                    ax.plot(quarter_list, pss_levels, label="Estrés", color=TutorDarkMoodPalette.URGENT, marker='o', linestyle='-', linewidth=2)
                    ax.set_title(f"Niveles por Cuatrimestre - {user_name}", fontsize=12, fontweight='bold', pad=10)
                    ax.set_xlabel("Cuatrimestre", fontsize=10)
                    ax.set_ylabel("Nivel", fontsize=10)
                    ax.set_ylim(0, 3)
                    ax.set_yticks([0, 1, 2, 3])
                    ax.set_yticklabels(['Bajo', 'Leve', 'Moderado', 'Alto'], fontsize=9)
                    ax.legend(loc='upper left', fontsize=9)
                    ax.grid(True, linestyle='--', alpha=0.7)
                    plt.xticks(rotation=45, fontsize=9)
                    plt.tight_layout()
                    temp_image = f"temp_{user_id}.png"
                    plt.savefig(temp_image, bbox_inches='tight', dpi=150)
                    plt.close()
                    doc.add_heading("Gráfico de Niveles", level=2)
                    doc.add_picture(temp_image, width=Inches(5.5))
                    if os.path.exists(temp_image):
                        os.remove(temp_image)

                # Recommendations
                recommendations = self.cache.get_user_recommendations(user_id) if self.cache else {}
                doc.add_heading("Recomendaciones", level=2)
                rec_map = {"BAI": "Ansiedad", "BDI": "Depresión", "PSS": "Estrés"}
                for questionnaire, text in recommendations.items():
                    name = rec_map.get(questionnaire, questionnaire)
                    doc.add_paragraph(f"{name}:", style='List Bullet')
                    p = doc.add_paragraph(text if text != "N/A" else "No hay recomendación disponible")
                    p.paragraph_format.left_indent = Inches(0.5)
                    for run in p.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                if not responses:
                    doc.add_paragraph("No hay datos de cuestionarios disponibles.")
                doc.add_page_break()

            # Save document to memory for download
            output_stream = io.BytesIO()
            doc.save(output_stream)
            output_stream.seek(0)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Reporte_{group}_{timestamp}.docx"
            base64_content = base64.b64encode(output_stream.getvalue()).decode('utf-8')
            data_url = f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64_content}"

            # Create a download link
            download_button = TextButton(
                content=Text(f"Descargar {filename}", color=TutorDarkMoodPalette.TEXT_ON_ACCENT),
                url=data_url,
                style=ButtonStyle(
                    bgcolor=TutorDarkMoodPalette.TRUST,
                    padding=padding.symmetric(vertical=10, horizontal=20),
                    shape=RoundedRectangleBorder(radius=8)
                )
            )
            self.page.overlay.append(download_button)
            self.show_snackbar(f"Reporte listo para descargar: {filename}", TutorDarkMoodPalette.SUCCESS_FEEDBACK)
            self.page.update()
            logger.info(f"[REPORT] Reporte generado para grupo {group} como enlace de descarga")

        except Exception as ex:
            self.show_snackbar(f"Error al generar reporte: {str(ex)}", TutorDarkMoodPalette.ERROR_FEEDBACK)
            logger.error(f"[REPORT ERROR] Error al generar reporte para grupo {group}: {str(ex)}")

    async def create_chart(self, student_id, cuatrimestre):
        if not student_id or not self.cache:
            return Text("No hay datos disponibles", color=TutorDarkMoodPalette.TEXT_MAIN)
        responses = self.cache.get_user_responses(student_id)
        if not responses:
            return Text("No hay respuestas registradas", color=TutorDarkMoodPalette.TEXT_MAIN)
        filtered_responses = responses
        if cuatrimestre != f"Todo {self.year}":
            start_date, end_date = self.get_cuatrimestre_dates(cuatrimestre)
            filtered_responses = [
                r for r in responses
                if self.is_response_in_date_range(r, start_date, end_date)
            ]
            if not filtered_responses:
                return Text(f"No hay datos para el período {cuatrimestre}",
                            color=TutorDarkMoodPalette.TEXT_MAIN)
        questionnaire_data = {"BAI": {}, "BDI": {}, "PSS": {}}
        month_year_data = {}
        for r in filtered_responses:
            date = r.get("date")
            if isinstance(date, str):
                try:
                    date = datetime.fromisoformat(date.replace("Z", "+00:00"))
                except ValueError:
                    continue
            q = r.get("questionnaire", "").strip().upper()
            if q not in questionnaire_data:
                continue
            level = r.get("level", 0)
            month_year = (date.year, date.month)
            if month_year not in month_year_data:
                month_year_data[month_year] = {}
            if q not in month_year_data[month_year] or date > month_year_data[month_year].get(q, {}).get("date", datetime.min.replace(tzinfo=timezone.utc)):
                month_year_data[month_year][q] = {"date": date, "level": level}
        for (year, month), q_data in month_year_data.items():
            month_str = self.get_month_name(month, year)
            for q in questionnaire_data:
                if q in q_data:
                    questionnaire_data[q][month_str] = q_data[q]["level"]
        data_series = []
        chart_colors = {
            "BAI": TutorDarkMoodPalette.ALERT,
            "BDI": TutorDarkMoodPalette.TRUST,
            "PSS": TutorDarkMoodPalette.URGENT
        }
        level_labels = {0: "Bajo", 1: "Leve", 2: "Moderado", 3: "Alto"}
        all_months = sorted(
            set([m for data in questionnaire_data.values() for m in data.keys()]),
            key=lambda m: (int(m.split()[1]), self.month_to_number(m.split()[0]))
        )
        for q_type in ["BAI", "BDI", "PSS"]:
            points = [
                LineChartDataPoint(
                    x=i+1,
                    y=questionnaire_data[q_type][month],
                    tooltip=f"{q_type}: {level_labels.get(questionnaire_data[q_type][month], 'N/A')} ({questionnaire_data[q_type][month]})"
                )
                for i, month in enumerate(all_months)
                if month in questionnaire_data[q_type]
            ]
            if points:
                data_series.append(
                    LineChartData(
                        data_points=points,
                        color=chart_colors[q_type],
                        curved=True,
                        stroke_width=2,
                        point=True,
                        below_line_bgcolor=Colors.with_opacity(0.15, chart_colors[q_type])
                    )
                )
        if not data_series:
            return Text("No hay datos válidos para graficar",
                        color=TutorDarkMoodPalette.TEXT_MAIN)
        chart = LineChart(
            data_series=data_series,
            horizontal_grid_lines=ChartGridLines(
                interval=1, color=Colors.with_opacity(0.2, Colors.ON_SURFACE), width=1
            ),
            vertical_grid_lines=ChartGridLines(
                interval=1, color=Colors.with_opacity(0.2, Colors.ON_SURFACE), width=1
            ),
            bottom_axis=ChartAxis(
                labels=[
                    ChartAxisLabel(
                        value=i+1,
                        label=Text(month.split()[0], size=10, color=TutorDarkMoodPalette.TEXT_MAIN)
                    )
                    for i, month in enumerate(all_months)
                ],
                labels_size=40
            ),
            left_axis=ChartAxis(
                labels=[
                    ChartAxisLabel(value=0, label=Text("0", color=TutorDarkMoodPalette.TEXT_MAIN)),
                    ChartAxisLabel(value=1, label=Text("1", color=TutorDarkMoodPalette.TEXT_MAIN)),
                    ChartAxisLabel(value=2, label=Text("2", color=TutorDarkMoodPalette.TEXT_MAIN)),
                    ChartAxisLabel(value=3, label=Text("3", color=TutorDarkMoodPalette.TEXT_MAIN))
                ],
                labels_size=40
            ),
            min_y=0,
            max_y=3,
            expand=True,
            tooltip_bgcolor=TutorDarkMoodPalette.CARD_SURFACE,
        )
        tables = []
        table_map = [
            ("Ansiedad", "BAI"),
            ("Depresión", "BDI"),
            ("Estrés", "PSS")
        ]
        for title, q in table_map:
            q_responses = sorted(
                [r for r in filtered_responses if r.get("questionnaire", "").strip().upper() == q],
                key=lambda r: r.get("date", datetime.min),
                reverse=True
            )
            table_rows = [
                DataRow(cells=[
                    DataCell(Text("Fecha", weight="bold", color=TutorDarkMoodPalette.TEXT_MAIN)),
                    DataCell(Text("Nivel", weight="bold", color=TutorDarkMoodPalette.TEXT_MAIN))
                ])
            ]
            if q_responses:
                for r in q_responses:
                    date = r.get("date")
                    if isinstance(date, str):
                        date = datetime.fromisoformat(date.replace("Z", "+00:00"))
                    date_str = date.strftime("%d %b")
                    level = r.get("level", "N/A")
                    table_rows.append(DataRow(cells=[
                        DataCell(Text(date_str, weight="bold", color=TutorDarkMoodPalette.TEXT_MAIN)),
                        DataCell(Text(str(level), weight="bold", color=TutorDarkMoodPalette.TEXT_MAIN))
                    ]))
            else:
                table_rows.append(DataRow(cells=[
                    DataCell(Text("No hay datos", color=TutorDarkMoodPalette.TEXT_MAIN)),
                    DataCell(Text(""))
                ]))
            table = DataTable(
                columns=[
                    DataColumn(label=Row([
                        Container(
                            width=10,
                            height=10,
                            bgcolor=chart_colors[q]
                        ),
                        Text(title, weight="bold", color=TutorDarkMoodPalette.TEXT_MAIN)
                    ], spacing=4, vertical_alignment=CrossAxisAlignment.CENTER)),
                    DataColumn(label=Text("")),
                ],
                rows=table_rows,
                border=border.all(1, TutorDarkMoodPalette.BORDER_SUBTLE),
                border_radius=8,
                expand=True
            )
            tables.append(
                Container(
                    content=Column([table], scroll="auto"),
                    height=200,
                    expand=True
                )
            )
        content = Column([
            chart,
            Text("Registros detallados", size=16, weight="bold", color=TutorDarkMoodPalette.TEXT_MAIN),
            Row(tables, spacing=10, expand=True)
        ], spacing=10, scroll="auto")
        return content

    def get_cuatrimestre_dates(self, cuatrimestre):
        year = self.year
        if cuatrimestre == f"Ene-Abr {year}":
            return datetime(year, 1, 1, tzinfo=timezone.utc), datetime(year, 4, 30, tzinfo=timezone.utc)
        elif cuatrimestre == f"May-Ago {year}":
            return datetime(year, 5, 1, tzinfo=timezone.utc), datetime(year, 8, 31, tzinfo=timezone.utc)
        elif cuatrimestre == f"Sep-Dic {year}":
            return datetime(year, 9, 1, tzinfo=timezone.utc), datetime(year, 12, 31, tzinfo=timezone.utc)
        return None, None

    def is_response_in_date_range(self, response, start_date, end_date):
        date = response.get("date")
        if not date:
            return False
        if isinstance(date, str):
            try:
                date = datetime.fromisoformat(date.replace("Z", "+00:00"))
            except ValueError:
                return False
        return start_date <= date <= end_date

    async def on_cuatrimestre_change(self, e):
        self.selected_cuatrimestre = e.control.value
        if self.selected_student:
            self.chart_container.content = await self.create_chart(
                self.selected_student.get('doc_id', ''),
                self.selected_cuatrimestre
            )
            if self.page:
                self.page.update()